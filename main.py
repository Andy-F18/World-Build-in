import tkinter as tk
import os
from tkinter import filedialog

import yaml
from docx import Document
from docx.shared import Cm

from Frames import CharPage as Chp
from Frames import LocationPage as Locp
from Frames import ItemsPage as Itp
from Frames import StoryPage as Stp


class WorldBuildIn:
    def __init__(self):
        # #################### BEGIN INIT VARS ####################
        self.size = {'h': 700, 'w': 1300}
        self.colors = {'bg1': '#ccc', 'bg2': "#555", 'menu1': '#555555', 'menu2': '#666666'}
        # #################### END INIT VARS ####################

        try:
            os.mkdir("Projects")
        except FileExistsError:
            pass

        # #################### BEGIN WINDOW ####################
        self.__root = tk.Tk()
        self.__root.title("World Build'in")
        self.__root.columnconfigure(0, weight=1)

        self.__root.configure(background=self.colors['bg1'])
        p_r = int(self.__root.winfo_screenwidth() / 2 - self.size['w'] / 2)
        p_d = int(self.__root.winfo_screenheight() / 2 - self.size['h'] / 2 - 50)
        self.__root.geometry("{}x{}+{}+{}".format(self.size['w'], self.size['h'], p_r, p_d))
        # #################### END WINDOW ####################

        self.__menu_p()
        self.__menu_h()
        self.__currentFrame = tk.Frame(self.__root)
        # self.__currentFrame.pack()

        self.worldName = tk.StringVar()
        self.workDir = tk.StringVar()
        self.workDir.set(os.path.abspath(".") + '\\Projects')
        self.defaultPath = os.path.abspath(".") + '\\Projects'

        self.__root.mainloop()

    def __menu_p(self):
        menu = tk.Menu(self.__root)

        fileMenu = tk.Menu(menu, tearoff=0, activebackground=self.colors['menu1'])
        fileMenu.add_command(label='New world', command=self.__new)
        fileMenu.add_command(label='Open world', command=self.__open)
        # fileMenu.add_command(label='Save')
        # fileMenu.add_command(label='Save as')
        fileMenu.add_command(label='Export to Docx', command=self.exportDocx)
        fileMenu.add_separator()
        fileMenu.add_command(label='Exit', command=self.__root.quit)
        menu.add_cascade(label='File', menu=fileMenu)

        addMenu = tk.Menu(menu, tearoff=0, activebackground=self.colors['menu1'])
        addMenu.add_command(label='Character', command=self.__addChar)
        addMenu.add_command(label='Location', command=self.__addLoc)
        addMenu.add_command(label='Item', command=self.__addItem)
        addMenu.add_command(label='Chapters', command=self.__addStory)
        menu.add_cascade(label='Add', menu=addMenu)

        self.__root.config(menu=menu)

    def __menu_h(self):
        f_menu = tk.Frame(self.__root, background=self.colors['menu1'])
        f_menu.columnconfigure(0, weight=1)
        f_menu.columnconfigure(1, weight=1)
        f_menu.columnconfigure(2, weight=1)
        f_menu.columnconfigure(3, weight=1)

        self.b_chara = tk.Button(f_menu, text="Character", relief='flat', command=self.__charPage,
                                 background=self.colors['menu1'], activebackground=self.colors['menu2'])
        self.b_place = tk.Button(f_menu, text="Locations", relief='flat', command=self.__locPage,
                                 background=self.colors['menu1'], activebackground=self.colors['menu2'])
        self.b_items = tk.Button(f_menu, text="Items", relief='flat', command=self.__itemPage,
                                 background=self.colors['menu1'], activebackground=self.colors['menu2'])
        self.b_story = tk.Button(f_menu, text="Story", relief='flat', command=self.__storyPage,
                                 background=self.colors['menu1'], activebackground=self.colors['menu2'])

        self.__import_file = tk.StringVar()
        self.b_chara.grid(column=0, row=0, sticky=tk.EW)
        self.b_place.grid(column=1, row=0, sticky=tk.EW)
        self.b_items.grid(column=2, row=0, sticky=tk.EW)
        self.b_story.grid(column=3, row=0, sticky=tk.EW)

        f_menu.grid(column=0, row=0, sticky=tk.EW)

    def __new(self):
        new = tk.Toplevel()
        new.title("New World")
        w = 500
        h = 150
        new.configure(background=self.colors["bg1"])
        pR = int(new.winfo_screenwidth() / 2 - w / 2)
        pD = int(new.winfo_screenheight() / 2 - h / 2 - 100)
        new.geometry("{}x{}+{}+{}".format(w, h, pR, pD))

        f_new = tk.Frame(new, background=self.colors['bg1'])
        # #################### BEGIN WORLD NAME FRAME ####################
        f_name = tk.Frame(f_new, background=self.colors['bg1'])
        tk.Label(f_name, text='Name :', background=self.colors['bg1']).grid(column=0, row=0)
        tk.Entry(f_name, textvariable=self.worldName, width=25).grid(column=1, row=0, padx=5)
        f_name.pack()
        # #################### END WORLD NAME FRAME ####################

        # #################### BEGIN WORk PATH FRAME ####################
        f_path = tk.Frame(f_new, background=self.colors['bg1'])
        tk.Label(f_path, text='Path :', background=self.colors['bg1']).grid(column=0, row=0)
        tk.Entry(f_path, textvariable=self.workDir, width=50).grid(column=1, row=0, padx=5)
        tk.Button(f_path, text='...', command=self.__browseFiles).grid(column=2, row=0)
        f_path.pack(pady=10)
        # #################### END WORK PATH FRAME ####################
        tk.Button(f_new, text='OK', command=new.destroy).pack()
        f_new.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

        new.grab_set()
        self.__root.wait_window(new)
        if self.worldName.get() == '':
            return

        os.mkdir(self.workDir.get() + "/" + self.worldName.get())
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Characters")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Characters/Images")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Locations")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Locations/Images")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Items")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Items/Images")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Story")
        # os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Images")
        file = open(self.workDir.get() + "/" + self.worldName.get() + "/." + self.worldName.get().replace(' ', '_', 99),
                    mode='w')
        file.write('name=' + self.worldName.get() + '\nOK=True')
        file.close()
        self.workDir.set(self.workDir.get() + '/' + self.worldName.get())

    def __open(self):
        self.__browseFiles()
        if os.listdir(self.workDir.get()).index('.' + self.worldName.get().replace(' ', '_', 999)) < 0:
            return

        file = open(self.workDir.get() + "/." + self.worldName.get().replace(' ', '_'), mode='r')
        fileTxt = file.readlines()
        file.close()
        if fileTxt[0] != 'name=' + self.worldName.get() + '\n':
            self.worldName.set('')
            self.workDir.set(self.defaultPath)
            return
        if fileTxt[1] != 'OK=True':
            self.worldName.set('')
            self.workDir.set(self.defaultPath)
            return

    def __addChar(self):
        if self.worldName.get() == '':
            return
        self.b_chara.config(background=self.colors['bg1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['menu1'])

        self.__currentFrame.destroy()
        frame = tk.Frame(self.__root, background=self.colors['bg1'])

        charP = Chp.CharPage(frame, self.colors, self.workDir)
        charP.createPage()

        self.__currentFrame = frame

    def __charPage(self):
        if self.worldName.get() == '':
            return
        self.b_chara.config(background=self.colors['bg1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['menu1'])

        self.__currentFrame.destroy()
        frame = tk.Frame(self.__root, background=self.colors['bg1'])

        charlist = os.listdir(self.workDir.get() + "/Characters")
        charlist.remove('Images')

        charP = Chp.CharPage(frame, self.colors, self.workDir)

        if len(charlist) != 0:
            charP.see(charlist[0])
        else:
            charP.createPage()

        self.__currentFrame = frame

    def __addLoc(self):
        if self.worldName.get() == '':
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['bg1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['menu1'])

        self.__currentFrame.destroy()
        placeFrame = tk.Frame(self.__root, background=self.colors['bg1'])

        locList = os.listdir(self.workDir.get() + '/Locations')

        locP = Locp.LocationPage(placeFrame, self.colors, self.workDir)

        locP.createPage()

        self.__currentFrame = placeFrame

    def __locPage(self):
        if self.worldName.get() == '':
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['bg1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['menu1'])

        self.__currentFrame.destroy()
        placeFrame = tk.Frame(self.__root, background=self.colors['bg1'])

        locList = os.listdir(self.workDir.get() + '/Locations')
        locList.remove('Images')

        locP = Locp.LocationPage(placeFrame, self.colors, self.workDir)

        if len(locList) != 0:
            locP.see(locList[0])
        else:
            locP.createPage()

        self.__currentFrame = placeFrame

    def __addItem(self):
        if self.worldName.get() == '':
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['bg1'])
        self.b_story.config(background=self.colors['menu1'])

        self.__currentFrame.destroy()
        itemFrame = tk.Frame(self.__root, background=self.colors['bg1'])

        locList = os.listdir(self.workDir.get() + '/Items')

        itP = Itp.ItemPage(itemFrame, self.colors, self.workDir)

        itP.createPage()

        self.__currentFrame = itemFrame

    def __itemPage(self):
        if self.worldName.get() == '':
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['bg1'])
        self.b_story.config(background=self.colors['menu1'])

        self.__currentFrame.destroy()
        itemFrame = tk.Frame(self.__root, background=self.colors['bg1'])

        locList = os.listdir(self.workDir.get() + '/Items')
        locList.remove('Images')

        itP = Itp.ItemPage(itemFrame, self.colors, self.workDir)

        if len(locList) != 0:
            itP.see(locList[0])
        else:
            itP.createPage()
        self.__currentFrame = itemFrame

    def __addStory(self):
        if self.worldName.get() == '':
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['bg1'])

        self.__currentFrame.destroy()
        storyFrame = tk.Frame(self.__root, background=self.colors['bg1'])

        storyP = Stp.StoryPage(storyFrame, self.colors, self.workDir)
        storyP.createPage()

        self.__currentFrame = storyFrame

    def __storyPage(self):
        if self.worldName.get() == '':
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['bg1'])

        self.__currentFrame.destroy()
        storyFrame = tk.Frame(self.__root, background=self.colors['bg1'])

        listStory = os.listdir(self.workDir.get()+'/Story')
        # listStory.remove('Images')
        storyP = Stp.StoryPage(storyFrame, self.colors, self.workDir)
        if len(listStory) != 0:
            storyP.see(listStory[0])
        else:
            storyP.createPage()

        self.__currentFrame = storyFrame

    def __browseFiles(self):
        rep = os.path.abspath(os.getcwd())
        filename_s = tk.filedialog.askdirectory(initialdir=rep,
                                                title="Select a Rep")
        if filename_s != "":
            filename = filename_s
        else:
            filename = self.workDir.get()
        # Change label contents
        self.workDir.set(filename)
        path = self.workDir.get().split('/')
        self.worldName.set(path[len(path) - 1])

    def exportDocx(self):
        if self.worldName.get() == '':
            return

        doc = Document()
        doc.add_heading(self.worldName.get(), 0)
        doc.add_page_break()

        # #################### BEGIN CHARACTERS ####################
        doc.add_heading('Characters', 1)
        charList = os.listdir(self.workDir.get()+'/Characters')
        charList.remove('Images')

        for char in charList:
            file = open(self.workDir.get() + '/Characters/' + char, mode='r')
            c = yaml.load(file, yaml.FullLoader)
            file.close()

            doc.add_heading(c['name'], 2)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            table.cell(0, 0).width = Cm(4)
            if c['photoFile']:
                picCell = table.rows[0].cells[0].paragraphs[0]
                run = picCell.add_run()
                run.add_picture(self.workDir.get()+'/Characters/Images/'+char.replace('.yml', '_FULL.png'),
                                width=Cm(3), height=Cm(3))

            p = table.rows[0].cells[1].add_paragraph()
            p.add_run('Age:').underline = True
            p.add_run(f"\t\t{c['age']}\n")
            p.add_run('Gender:').underline = True
            p.add_run(f"\t{c['gender']}")

            doc.add_heading('About', 3)
            doc.add_paragraph(c['about'])

            doc.add_page_break()
        # #################### END CHARACTERS ####################

        # #################### BEGIN LOCATIONS ####################
        doc.add_heading('Locations', 1)
        locList = os.listdir(self.workDir.get()+"/Locations")
        locList.remove('Images')

        for loc in locList:
            file = open(self.workDir.get()+'/locations/'+loc, mode='r')
            l = yaml.load(file, yaml.FullLoader)
            file.close()

            doc.add_heading(l['name'], 2)
            if l['photoFile']:
                doc.add_picture(self.workDir.get() + '/Locations/Images/' + loc.replace('.yml', '_FULL.png'),
                                width=Cm(15))

            if l['location'] != '':
                p = doc.add_paragraph()
                p.add_run('Location:').underline = True
                p.add_run(f"\t{l['location']}")

            doc.add_heading('Places', 3)
            if len(l['places']) != 0:
                p = doc.add_paragraph(style='List Bullet')
                for place in l['places']:
                    p.add_run(place)
            else:
                doc.add_paragraph('No places here.')

            doc.add_heading('About', 3)
            doc.add_paragraph(l['about'])

            doc.add_page_break()
        # #################### END LOCATIONS ####################

        # #################### BEGIN ITEMS ####################
        doc.add_heading('Items', 1)
        itemList = os.listdir(self.workDir.get()+'/Items')
        itemList.remove('Images')

        for item in itemList:
            file = open(self.workDir.get() + '/Items/' + item, mode='r')
            i = yaml.load(file, yaml.FullLoader)
            file.close()

            doc.add_heading(i['name'], 2)
            if i['photoFile']:
                doc.add_picture(self.workDir.get()+'/Items/Images/'+item.replace('.yml', '_FULL.png'),
                                height=Cm(3))

            if i['location'] != '':
                p = doc.add_paragraph()
                p.add_run('Location:').underline = True
                p.add_run(f"\t{i['location']}")

            doc.add_heading('About', 3)
            doc.add_paragraph(i['about'])

            doc.add_page_break()
        # #################### END ITEMS ####################

        # #################### BEGIN Story ####################
        doc.add_heading('Story', 1)
        chapList = os.listdir(self.workDir.get()+'/Story')

        for chap in chapList:
            file = open(self.workDir.get() + '/Story/' + chap, mode='r')
            c = yaml.load(file, yaml.FullLoader)
            file.close()

            doc.add_heading(c['name'], 2)
            if c['location'] != '':
                p = doc.add_paragraph()
                p.add_run('location:').underline = True
                p.add_run(f"\t{c['location']}")

            doc.add_heading('Content', 3)
            doc.add_paragraph(c['about'])

            doc.add_page_break()
        # #################### END Story ####################

        doc.save(self.workDir.get()+"/"+self.worldName.get().replace(' ', '_', 99)+'.docx')


if __name__ == '__main__':
    WorldBuildIn()
