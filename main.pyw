import tkinter as tk
import os
from tkinter import filedialog
from tkinter import messagebox

import yaml
from docx import Document
from docx.shared import Cm

from Frames import CharPage as ChPage
from Frames import LocationPage as LocPage
from Frames import ItemsPage as ItPage
from Frames import StoryPage as StPage


class WorldBuildIn:
    def __init__(self):
        # #################### BEGIN INIT VARS ####################
        self.size = {'h': 700, 'w': 1300}
        self.colors = {'bg1': '#ccc', 'bg2': "#555", 'menu1': '#555555', 'menu2': '#666666'}
        # #################### END INIT VARS ####################

        try:
            os.mkdir("Projects")
        except FileExistsError:
            pass

        # #################### BEGIN WINDOW ####################
        self.__root = tk.Tk()
        self.__root.title("World Build")
        self.__root.columnconfigure(0, weight=1)

        self.__root.configure(background=self.colors['bg1'])
        p_r = int(self.__root.winfo_screenwidth() / 2 - self.size['w'] / 2)
        p_d = int(self.__root.winfo_screenheight() / 2 - self.size['h'] / 2 - 50)
        self.__root.geometry("{}x{}+{}+{}".format(self.size['w'], self.size['h'], p_r, p_d))
        # #################### END WINDOW ####################

        self.__menu_p()
        self.__menu_h()
        self.__currentFrame = None
        # self.__currentFrame.pack()

        self.worldName = tk.StringVar()
        self.workDir = tk.StringVar()
        self.workDir.set(os.path.abspath(".") + '\\Projects')
        self.defaultPath = os.path.abspath(".") + '\\Projects'

        self.__root.mainloop()

    def __menu_p(self):
        menu = tk.Menu(self.__root)

        fileMenu = tk.Menu(menu, tearoff=0, activebackground=self.colors['menu1'])
        fileMenu.add_command(label='New world', command=self.__new)
        fileMenu.add_command(label='Open world', command=self.__open)
        # fileMenu.add_command(label='Save', command=self.message)
        # fileMenu.add_command(label='Save as')
        fileMenu.add_command(label='Export to Docx', command=self.exportDocx)
        fileMenu.add_separator()
        fileMenu.add_command(label='Exit', command=self.__root.quit)
        menu.add_cascade(label='File', menu=fileMenu)

        addMenu = tk.Menu(menu, tearoff=0, activebackground=self.colors['menu1'])
        addMenu.add_command(label='Character', command=self.__addChar)
        addMenu.add_command(label='Location', command=self.__addLoc)
        addMenu.add_command(label='Item', command=self.__addItem)
        addMenu.add_command(label='Chapters', command=self.__addStory)
        menu.add_cascade(label='Add', menu=addMenu)

        self.__root.config(menu=menu)

    def __menu_h(self):
        f_menu = tk.Frame(self.__root, background=self.colors['menu1'])
        f_menu.columnconfigure(0, weight=1)
        f_menu.columnconfigure(1, weight=1)
        f_menu.columnconfigure(2, weight=1)
        f_menu.columnconfigure(3, weight=1)

        self.b_chara = tk.Button(f_menu, text="Character", relief='flat', command=self.charPage,
                                 background=self.colors['menu1'], activebackground=self.colors['menu2'])
        self.b_place = tk.Button(f_menu, text="Locations", relief='flat', command=self.locPage,
                                 background=self.colors['menu1'], activebackground=self.colors['menu2'])
        self.b_items = tk.Button(f_menu, text="Items", relief='flat', command=self.itemPage,
                                 background=self.colors['menu1'], activebackground=self.colors['menu2'])
        self.b_story = tk.Button(f_menu, text="Story", relief='flat', command=self.storyPage,
                                 background=self.colors['menu1'], activebackground=self.colors['menu2'])

        self.__import_file = tk.StringVar()
        self.b_chara.grid(column=0, row=0, sticky=tk.EW)
        self.b_place.grid(column=1, row=0, sticky=tk.EW)
        self.b_items.grid(column=2, row=0, sticky=tk.EW)
        self.b_story.grid(column=3, row=0, sticky=tk.EW)

        f_menu.grid(column=0, row=0, sticky=tk.EW)

    def __new(self):
        new = tk.Toplevel()
        new.title("New World")
        w = 500
        h = 150
        new.configure(background=self.colors["bg1"])
        pR = int(new.winfo_screenwidth() / 2 - w / 2)
        pD = int(new.winfo_screenheight() / 2 - h / 2 - 100)
        new.geometry("{}x{}+{}+{}".format(w, h, pR, pD))

        f_new = tk.Frame(new, background=self.colors['bg1'])
        # #################### BEGIN WORLD NAME FRAME ####################
        f_name = tk.Frame(f_new, background=self.colors['bg1'])
        tk.Label(f_name, text='Name :', background=self.colors['bg1']).grid(column=0, row=0)
        tk.Entry(f_name, textvariable=self.worldName, width=25).grid(column=1, row=0, padx=5)
        f_name.pack()
        # #################### END WORLD NAME FRAME ####################

        # #################### BEGIN WORk PATH FRAME ####################
        f_path = tk.Frame(f_new, background=self.colors['bg1'])
        tk.Label(f_path, text='Path :', background=self.colors['bg1']).grid(column=0, row=0)
        tk.Entry(f_path, textvariable=self.workDir, width=50).grid(column=1, row=0, padx=5)
        tk.Button(f_path, text='...', command=self.__browseFiles).grid(column=2, row=0)
        f_path.pack(pady=10)
        # #################### END WORK PATH FRAME ####################
        tk.Button(f_new, text='OK', command=new.destroy).pack()
        f_new.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

        new.grab_set()
        self.__root.wait_window(new)
        if self.worldName.get() == '':
            return

        os.mkdir(self.workDir.get() + "/" + self.worldName.get())
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Characters")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Characters/Images")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Locations")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Locations/Images")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Items")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Items/Images")
        os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Story")
        # os.mkdir(self.workDir.get() + "/" + self.worldName.get() + "/Images")
        file = open(self.workDir.get() + "/" + self.worldName.get() + "/." + self.worldName.get().replace(' ', '_', 99),
                    mode='w')
        file.write('name=' + self.worldName.get() + '\nOK=True')
        file.close()
        self.workDir.set(self.workDir.get() + '/' + self.worldName.get())
        self.charPage()

    def __open(self):
        self.__browseFiles()
        if os.listdir(self.workDir.get()).index('.' + self.worldName.get().replace(' ', '_', 999)) < 0:
            return

        file = open(self.workDir.get() + "/." + self.worldName.get().replace(' ', '_'), mode='r')
        fileTxt = file.readlines()
        file.close()
        if fileTxt[0] != 'name=' + self.worldName.get() + '\n':
            self.worldName.set('')
            self.workDir.set(self.defaultPath)
            return
        if fileTxt[1] != 'OK=True':
            self.worldName.set('')
            self.workDir.set(self.defaultPath)
            return

        self.charPage()

    def __addChar(self):
        if self.worldName.get() == '':
            self.warning('Add character', 'Open/Create a world before.')
            return
        self.b_chara.config(background=self.colors['bg1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['menu1'])

        charP = ChPage.CharPage(self, self.chFrame(), self.colors, self.workDir)
        charP.createPage()

        self.__currentFrame = charP

    def charPage(self, char=None):
        if self.worldName.get() == '':
            self.warning('Character', 'Open/Create a world before.')
            return
        self.b_chara.config(background=self.colors['bg1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['menu1'])

        charList = os.listdir(self.workDir.get() + "/Characters")
        charList.remove('Images')

        charP = ChPage.CharPage(self, self.chFrame(), self.colors, self.workDir)

        if len(charList) != 0:
            if char is None:
                charP.see(charList[0])
            else:
                charP.see(char)
        else:
            charP.createPage()

        self.__currentFrame = charP

    def __addLoc(self):
        if self.worldName.get() == '':
            self.warning('Add location', 'Open/Create a world before.')
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['bg1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['menu1'])

        locP = LocPage.LocationPage(self, self.chFrame(), self.colors, self.workDir)

        locP.createPage()

        self.__currentFrame = locP

    def locPage(self, loc=None):
        if self.worldName.get() == '':
            self.warning('Locations', 'Open/Create a world before.')
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['bg1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['menu1'])

        locList = os.listdir(self.workDir.get() + '/Locations')
        locList.remove('Images')

        locP = LocPage.LocationPage(self, self.chFrame(), self.colors, self.workDir)

        if len(locList) != 0:
            if loc is None:
                locP.see(locList[0])
            else:
                locP.see(loc)
        else:
            locP.createPage()

        self.__currentFrame = locP

    def __addItem(self):
        if self.worldName.get() == '':
            self.warning('Add item', 'Open/Create a world before.')
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['bg1'])
        self.b_story.config(background=self.colors['menu1'])

        itP = ItPage.ItemPage(self, self.chFrame(), self.colors, self.workDir)

        itP.createPage()

        self.__currentFrame = itP

    def itemPage(self, item=None):
        if self.worldName.get() == '':
            self.warning('Items', 'Open/Create a world before.')
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['bg1'])
        self.b_story.config(background=self.colors['menu1'])

        locList = os.listdir(self.workDir.get() + '/Items')
        locList.remove('Images')

        itP = ItPage.ItemPage(self, self.chFrame(), self.colors, self.workDir)

        if len(locList) != 0:
            if item is None:
                itP.see(locList[0])
            else:
                itP.see(item)
        else:
            itP.createPage()
        self.__currentFrame = itP

    def __addStory(self):
        if self.worldName.get() == '':
            self.warning('Add chapter', 'Open/Create a world before.')
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['bg1'])

        storyP = StPage.StoryPage(self, self.chFrame(), self.colors, self.workDir)
        storyP.createPage()

        self.__currentFrame = storyP

    def storyPage(self, story=None):
        if self.worldName.get() == '':
            self.warning('Story', 'Open/Create a world before.')
            return
        self.b_chara.config(background=self.colors['menu1'])
        self.b_place.config(background=self.colors['menu1'])
        self.b_items.config(background=self.colors['menu1'])
        self.b_story.config(background=self.colors['bg1'])

        listStory = os.listdir(self.workDir.get()+'/Story')
        # listStory.remove('Images')
        storyP = StPage.StoryPage(self, self.chFrame(), self.colors, self.workDir)
        if len(listStory) != 0:
            if story is None:
                storyP.see(listStory[0])
            else:
                storyP.see(story)
        else:
            storyP.createPage()

        self.__currentFrame = storyP

    def __browseFiles(self):
        rep = os.path.abspath(os.getcwd())
        filename_s = tk.filedialog.askdirectory(initialdir=rep,
                                                title="Select a Rep")
        if filename_s != "":
            filename = filename_s
        else:
            filename = self.workDir.get()
        # Change label contents
        self.workDir.set(filename)
        path = self.workDir.get().split('/')
        self.worldName.set(path[len(path) - 1])

    def exportDocx(self):
        if self.worldName.get() == '':
            self.warning('Export to docx', 'Open/Create a world before.')
            return

        doc = Document()
        doc.add_heading(self.worldName.get(), 0)
        doc.add_page_break()

        # #################### BEGIN CHARACTERS ####################
        doc.add_heading('Characters', 1)
        charList = os.listdir(self.workDir.get()+'/Characters')
        charList.remove('Images')

        for char in charList:
            file = open(self.workDir.get() + '/Characters/' + char, mode='r')
            c = yaml.load(file, yaml.FullLoader)
            file.close()

            doc.add_heading(c['name'], 2)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            table.cell(0, 0).width = Cm(4)
            if c['photoFile']:
                picCell = table.cell(0, 0).paragraphs[0]
                run = picCell.add_run()
                run.add_picture(self.workDir.get()+'/Characters/Images/'+char.replace('.yml', '_FULL.png'),
                                width=Cm(3), height=Cm(3))

            p = table.cell(0, 1).add_paragraph()
            p.add_run('Age:').underline = True
            p.add_run(f"\t\t{c['age']}\n")
            p.add_run('Gender:').underline = True
            p.add_run(f"\t{c['gender']}")

            doc.add_heading('About', 3)
            doc.add_paragraph(c['about'])

            doc.add_page_break()
        # #################### END CHARACTERS ####################

        # #################### BEGIN LOCATIONS ####################
        doc.add_heading('Locations', 1)
        locList = os.listdir(self.workDir.get()+"/Locations")
        locList.remove('Images')

        for loc in locList:
            file = open(self.workDir.get()+'/locations/'+loc, mode='r')
            l = yaml.load(file, yaml.FullLoader)
            file.close()

            doc.add_heading(l['name'], 2)
            if l['photoFile']:
                doc.add_picture(self.workDir.get() + '/Locations/Images/' + loc.replace('.yml', '_FULL.png'),
                                width=Cm(15))

            if l['location'] != '':
                p = doc.add_paragraph()
                p.add_run('Location:').underline = True
                p.add_run(f"\t{l['location']}")

            doc.add_heading('Places', 3)
            if len(l['places']) != 0:
                for place in l['places']:
                    doc.add_paragraph(place, style='List Bullet')
            else:
                doc.add_paragraph('No places here.')

            doc.add_heading('About', 3)
            doc.add_paragraph(l['about'])

            doc.add_page_break()
        # #################### END LOCATIONS ####################

        # #################### BEGIN ITEMS ####################
        doc.add_heading('Items', 1)
        itemList = os.listdir(self.workDir.get()+'/Items')
        itemList.remove('Images')

        for item in itemList:
            file = open(self.workDir.get() + '/Items/' + item, mode='r')
            i = yaml.load(file, yaml.FullLoader)
            file.close()

            doc.add_heading(i['name'], 2)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            table.cell(0, 0).width = Cm(4)
            if i['photoFile']:
                picCell = table.cell(0, 0).paragraphs[0]
                run = picCell.add_run()
                run.add_picture(self.workDir.get()+'/Items/Images/'+item.replace('.yml', '_FULL.png'),
                                height=Cm(3))

            if i['location'] != '':
                p = table.cell(0, 1).add_paragraph()
                p.add_run('Location:').underline = True
                p.add_run(f"\t{i['location']}")

            doc.add_heading('About', 3)
            doc.add_paragraph(i['about'])

            doc.add_page_break()
        # #################### END ITEMS ####################

        # #################### BEGIN Story ####################
        doc.add_heading('Story', 1)
        chapList = os.listdir(self.workDir.get()+'/Story')

        for chap in chapList:
            file = open(self.workDir.get() + '/Story/' + chap, mode='r')
            c = yaml.load(file, yaml.FullLoader)
            file.close()

            doc.add_heading(c['name'], 2)
            if c['location'] != '':
                p = doc.add_paragraph()
                p.add_run('location:').underline = True
                p.add_run(f"\t{c['location']}")

            doc.add_heading('Content', 3)
            doc.add_paragraph(c['about'])

            doc.add_page_break()
        # #################### END Story ####################

        doc.save(self.workDir.get()+"/"+self.worldName.get().replace(' ', '_', 99)+'.docx')
        self.message('Exportation', 'Exportation successful!')

    def chFrame(self):
        try:
            self.__currentFrame.save()
        except AttributeError:
            pass

        try:
            self.__currentFrame.destroy()
        except AttributeError:
            pass

        return tk.Frame(self.__root, background=self.colors['bg1'])

    @staticmethod
    def warning(title, msg):
        messagebox.showwarning(title, msg)

    @staticmethod
    def message(title, msg):
        messagebox.showinfo(title, msg)


if __name__ == '__main__':
    WorldBuildIn()
